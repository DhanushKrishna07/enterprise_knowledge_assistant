"""
app/documents/loaders.py — Unified document loaders for PDF, DOCX, TXT/MD, and CSV.

Each loader returns:
  - A DocumentMeta for the document
  - A list of (page_number | None, text) tuples representing raw page text
  - A list of TableChunk objects (for PDF)
  - An extraction_method string per page
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.core.config import get_settings
from app.documents.hashing import document_id, file_checksum
from app.documents.metadata import DocumentMeta
from app.documents.ocr import ocr_pdf_page
from app.documents.pdf_parser import PageResult, extract_pdf
from app.documents.table_extractor import TableChunk, tables_from_pdfplumber_page

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Raw output from a loader — ready for chunking."""

    doc_meta: DocumentMeta
    pages: list[tuple[int | None, str]]  # (page_number_or_None, text)
    extraction_methods: list[str]  # parallel to pages
    table_chunks: list[TableChunk] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    ocr_page_count: int = 0


# ── Sidecar metadata loader ───────────────────────────────────────────────────


def _load_sidecar(path: Path) -> dict[str, Any]:
    """Load optional YAML sidecar file for custom metadata (same name + .meta.yaml)."""
    sidecar = path.with_suffix(".meta.yaml")
    if not sidecar.exists():
        sidecar = path.parent / (path.name + ".meta.yaml")
    if sidecar.exists():
        try:
            import yaml  # type: ignore[import]

            with open(sidecar, encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        except Exception as exc:
            logger.warning("Could not read sidecar %s: %s", sidecar, exc)
    return {}


def _doc_meta_from_path(
    path: Path, checksum: str, file_type: str, sidecar: dict[str, Any]
) -> DocumentMeta:
    doc_id = document_id(path.name, checksum)
    return DocumentMeta(
        document_id=doc_id,
        filename=path.name,
        file_type=file_type,
        checksum=checksum,
        department=sidecar.get("department", "general"),
        author=sidecar.get("author", ""),
        tags=sidecar.get("tags", []),
        policy_version=str(sidecar.get("policy_version", "")),
        upload_date=str(sidecar.get("upload_date", "")),
        access_roles=sidecar.get("access_roles", ["employee"]),
    )


# ── PDF Loader ────────────────────────────────────────────────────────────────


def load_pdf(path: Path) -> LoadedDocument:
    settings = get_settings()
    checksum = file_checksum(path)
    sidecar = _load_sidecar(path)
    doc_meta = _doc_meta_from_path(path, checksum, "pdf", sidecar)

    page_results: list[PageResult] = extract_pdf(
        path,
        parser=settings.pdf_parser,
        fallback_parser=settings.pdf_fallback_parser,
        ocr_min_chars=settings.ocr_min_text_chars_per_page,
        extract_tables=settings.enable_table_extraction,
    )

    pages: list[tuple[int | None, str]] = []
    methods: list[str] = []
    table_chunks: list[TableChunk] = []
    warnings: list[str] = []
    ocr_count = 0

    for pr in page_results:
        text = pr.text

        # OCR fallback for low-text pages
        if pr.needs_ocr and settings.enable_ocr:
            ocr_result = ocr_pdf_page(
                path,
                page_number=pr.page_number,
                engine=settings.ocr_engine,
            )
            if ocr_result.success and ocr_result.text.strip():
                text = ocr_result.text
                method = f"{settings.ocr_engine}_ocr"
                ocr_count += 1
            else:
                method = pr.extraction_method
                if ocr_result.error:
                    warnings.append(f"OCR failed on page {pr.page_number}: {ocr_result.error}")
        elif pr.needs_ocr and not settings.enable_ocr:
            warnings.append(
                f"Page {pr.page_number} has <{settings.ocr_min_text_chars_per_page} chars; OCR disabled."
            )
            method = pr.extraction_method
        else:
            method = pr.extraction_method

        pages.append((pr.page_number, text))
        methods.append(method)

        # Extract tables from pdfplumber raw data
        if pr.tables_raw and settings.enable_table_extraction:
            section = ""  # heading detection is done later in chunker
            tcs = tables_from_pdfplumber_page(pr.tables_raw, pr.page_number, section_title=section)
            table_chunks.extend(tcs)

    return LoadedDocument(
        doc_meta=doc_meta,
        pages=pages,
        extraction_methods=methods,
        table_chunks=table_chunks,
        warnings=warnings,
        ocr_page_count=ocr_count,
    )


# ── DOCX Loader ───────────────────────────────────────────────────────────────


def load_docx(path: Path) -> LoadedDocument:
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        logger.error("python-docx not installed; cannot load %s.", path)
        return _empty(path, "docx")

    checksum = file_checksum(path)
    sidecar = _load_sidecar(path)
    doc_meta = _doc_meta_from_path(path, checksum, "docx", sidecar)

    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.error("Failed to open DOCX %s: %s", path, exc)
        return _empty(path, "docx")

    paragraphs: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    table_chunks: list[TableChunk] = []
    for t_idx, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        from app.documents.table_extractor import _table_to_markdown

        md = _table_to_markdown(rows)
        cols = rows[0] if rows else []
        table_chunks.append(
            TableChunk(
                page_number=None,  # type: ignore[arg-type]
                table_index=t_idx,
                markdown=md,
                columns=cols,
                row_count=max(0, len(rows) - 1),
                engine="python-docx",
            )
        )

    return LoadedDocument(
        doc_meta=doc_meta,
        pages=[(None, full_text)],
        extraction_methods=["python-docx"],
        table_chunks=table_chunks,
    )


# ── TXT / MD Loader ───────────────────────────────────────────────────────────


def load_text(path: Path) -> LoadedDocument:
    checksum = file_checksum(path)
    sidecar = _load_sidecar(path)
    ext = path.suffix.lower().lstrip(".")
    doc_meta = _doc_meta_from_path(path, checksum, ext, sidecar)

    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.error("Cannot read %s: %s", path, exc)
        return _empty(path, ext)

    return LoadedDocument(
        doc_meta=doc_meta,
        pages=[(None, text)],
        extraction_methods=["plaintext"],
    )


# ── CSV Loader ────────────────────────────────────────────────────────────────


def load_csv(path: Path) -> LoadedDocument:
    try:
        import pandas as pd  # type: ignore[import]
    except ImportError:
        logger.error("pandas not installed; cannot load CSV %s.", path)
        return _empty(path, "csv")

    checksum = file_checksum(path)
    sidecar = _load_sidecar(path)
    doc_meta = _doc_meta_from_path(path, checksum, "csv", sidecar)

    try:
        df = pd.read_csv(str(path))
    except Exception as exc:
        logger.error("Cannot parse CSV %s: %s", path, exc)
        return _empty(path, "csv")

    # Convert each row to a text record
    rows_text: list[str] = []
    for _, row in df.iterrows():
        row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if str(val).strip())
        rows_text.append(row_text)

    full_text = "\n".join(rows_text)

    # Also expose as a table chunk
    rows_list = [list(df.columns)] + df.astype(str).values.tolist()
    from app.documents.table_extractor import _table_to_markdown

    md = _table_to_markdown(rows_list)
    tc = TableChunk(
        page_number=None,  # type: ignore[arg-type]
        table_index=0,
        markdown=md,
        columns=list(df.columns),
        row_count=len(df),
        engine="pandas",
    )

    return LoadedDocument(
        doc_meta=doc_meta,
        pages=[(None, full_text)],
        extraction_methods=["csv"],
        table_chunks=[tc],
    )


# ── Router ────────────────────────────────────────────────────────────────────


def load_document(path: str | Path) -> LoadedDocument | None:
    """Route to the correct loader based on file extension."""
    p = Path(path)
    if not p.exists():
        logger.error("File not found: %s", p)
        return None

    ext = p.suffix.lower()
    if ext == ".pdf":
        return load_pdf(p)
    elif ext in (".docx", ".doc"):
        return load_docx(p)
    elif ext in (".txt", ".md"):
        return load_text(p)
    elif ext == ".csv":
        return load_csv(p)
    else:
        logger.warning("Unsupported file type: %s", ext)
        return None


# ── Helpers ───────────────────────────────────────────────────────────────────


def _empty(path: Path, file_type: str) -> LoadedDocument:
    checksum = "0" * 64
    doc_meta = DocumentMeta(
        document_id=document_id(path.name, checksum),
        filename=path.name,
        file_type=file_type,
        checksum=checksum,
    )
    return LoadedDocument(
        doc_meta=doc_meta,
        pages=[],
        extraction_methods=[],
        warnings=[f"Failed to load {path.name}"],
    )
